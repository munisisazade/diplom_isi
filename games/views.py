import os

import math
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.shortcuts import render, redirect
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.views.generic import TemplateView, View
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.utils import timezone
from django.utils.timezone import localtime as loc
from social_django.models import UserSocialAuth

from games.models import Question, Answer, UserQuestionHistory, GameTime, LeaderBoard, WebsiteHeader, Repeating, \
    MonthBoard, WeekBoard, WeeklyResults, MonthlyResults, LeaderResults
from games.options.decorator import UserGameSessionIdCheck, CheckLeaderBorad, increase_point, FacebookFriendsList
from games.tasks import calc_weekly, calc_leaderbord, calc_monthly
from base_user.tools.common import game_session_generate, get_four_month_calendar, day_range_utc, current_week_range, current_month_range
from django.contrib.auth import get_user_model
from django.db.models import Q
from games.tasks import leaderbord_count_week_and_month
from django.urls import reverse
from django.http import HttpResponsePermanentRedirect
import calendar
# Create your views here.
import random
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Inches
# import the logging library
import logging

# Get an instance of a logger
logr = logging.getLogger(__name__)

User = get_user_model()


# class BaseView(View):
#
#     def get_context_data(self):
#         context = {}
#         context['base_statistic'] = GameTime.objects.filter()
#         return context

# class BaseView(View):
#     def get_context_data(self, **kwargs):
#         context = {}
#         all_user = GameTime.objects.filter(done=True)
#         context['winners'] = [user for user in all_user if user.get_right_answer_count > 7]
#         return context

class BaseIndexView(UserGameSessionIdCheck, TemplateView):
    template_name = 'production/home/index.html'

    def get_context_data(self, **kwargs):
        context = super(BaseIndexView, self).get_context_data(**kwargs)
        context['header'] = WebsiteHeader.objects.all().last()
        context['winners'] = LeaderBoard.objects.all().order_by('-score', 'duration')[:10]
        return context


class IndexView(View):
    template_name = 'production/test/index.html'

    def get(self, request):
        return JsonResponse({"name": [1, 2, 3, 4]})

        # def get_context_data(self, **kwargs):
        #     context = {}
        #     context['questions'] = Question.objects.all().order_by('?')[:10]
        #     return context


class OnlineTestView(TemplateView):
    template_name = 'production/test/index.html'
    limit_template = 'production/limit/index.html'

    def get(self, request, *args, **kwargs):
        try:
            if GameTime.objects.filter(game_session=request.user.game_session).last():
                game = GameTime.objects.filter(game_session=request.user.game_session).last()
                if not game.done:
                    c = loc(loc(timezone.now())) - game.game_time
                    duration = divmod(c.days * 86400 + c.seconds, 60)
                    game.duration = timezone.timedelta(minutes=duration[0], seconds=duration[1])
                    game.done = True
                    game.last_done = True
                    game.save()
                    game.player.game_session = game_session_generate()
                    game.player.save()
                return redirect("/")
            else:
                if not request.user.aditional_chance:
                    if GameTime.objects.filter(Q(player=request.user),
                                               Q(player__unlimited=False),
                                               Q(game_time__gte=day_range_utc()[0]),
                                               Q(game_time__lte=day_range_utc()[1])).count() < 3 or self.check_user_social_share_permission():
                        new_game = GameTime(
                            player=request.user,
                            game_session=request.user.game_session,
                            over=False
                        )
                        new_game.save()
                        return render(request, self.template_name, self.get_context_data())
                    else:
                        return render(request, self.limit_template, self.get_context_data())
                else:
                    if request.user.aditional_games_count != 0 and GameTime.objects.filter(Q(player=request.user),
                                                                                           Q(player__unlimited=False),
                                                                                           Q(game_time__gte=day_range_utc()[0]),
                                                                                           Q(game_time__lte=day_range_utc()[1])).count() >= 3:
                        request.user.aditional_games_count -= 1
                        if request.user.aditional_games_count == 0:
                            request.user.aditional_chance = False
                        request.user.save()
                        new_game = GameTime(
                            player=request.user,
                            game_session=request.user.game_session,
                            over=False
                        )
                        new_game.save()
                        return render(request, self.template_name, self.get_context_data())
                    else:
                        if GameTime.objects.filter(Q(player=request.user),
                                                   Q(player__unlimited=False),
                                                   Q(game_time__gte=day_range_utc()[0]),
                                                   Q(game_time__lte=day_range_utc()[1])).count() < 3 or self.check_user_social_share_permission():
                            new_game = GameTime(
                                player=request.user,
                                game_session=request.user.game_session,
                                over=False
                            )
                            new_game.save()
                            return render(request, self.template_name, self.get_context_data())
                        else:
                            return render(request, self.limit_template, self.get_context_data())

        except:
            return render(request, self.template_name, self.get_context_data())

    def day_range(self):
        current_time = timezone.now()
        if current_time.hour >= 20:
            day_start = timezone.datetime(current_time.year,
                                          current_time.month,
                                          current_time.day,
                                          20,
                                          0).replace(tzinfo=current_time.tzinfo)
            day_end = timezone.datetime(current_time.year,
                                          current_time.month,
                                          current_time.day+1,
                                          20,
                                          0).replace(tzinfo=current_time.tzinfo)
        else:
            day_start = timezone.datetime(current_time.year,
                                          current_time.month,
                                          current_time.day-1,
                                          20,
                                          0).replace(tzinfo=current_time.tzinfo)
            day_end = timezone.datetime(current_time.year,
                                        current_time.month,
                                        current_time.day,
                                        20,
                                        0).replace(tzinfo=current_time.tzinfo)
        return [day_start, day_end]

    def check_token(self, token):
        try:
            User.objects.get(game_session=token)
            return True
        except:
            return False

    def get_context_data(self, **kwargs):
        context = {}
        exclude_game = eval(self.request.user.exclude_question)
        if exclude_game:
            context['questions'] = Question.objects.all().order_by('?').exclude(id__in=list(set(exclude_game)))[:10]
        else:
            context['questions'] = Question.objects.all().order_by('?')[:10]
        if context['questions'].count() < 10:
            self.request.user.exclude_question = "[]"
            self.request.user.save()

        # context['months_winner'] = MonthBoard.objects.filter(month_name=str(current_month_range())).order_by('-score', 'duration')[:10]
        # context['weeks_winner'] = WeekBoard.objects.filter(week_name=str(current_week_range())).order_by('-score', 'duration')[:10]

        today = timezone.now() + timezone.timedelta(hours=4)
        start_month = 1
        end_month = calendar.monthrange(today.year, today.month)[1]

        start_month = timezone.datetime.strptime("{0}-{1}-{2}".format(today.year, today.month, start_month), '%Y-%m-%d')
        start_month = start_month.replace(tzinfo=today.tzinfo)

        end_month = timezone.datetime.strptime("{0}-{1}-{2}".format(today.year, today.month, end_month), '%Y-%m-%d')
        end_month = end_month.replace(tzinfo=today.tzinfo)

        monthly_data = MonthlyResults.objects.filter(
            start=start_month.date(),
            end=end_month.date()
        ).last()
        monthly_result = eval(monthly_data.result) if monthly_data else []
        monthly_users = User.objects.filter(id__in=[p['player'] for p in monthly_result])
        context['months_winner'] = [
            {"user": monthly_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in monthly_result]

        # start_week = timezone.datetime(today.year, today.month - 1, 25)  # default
        # end_week = timezone.datetime(today.year, today.month, 8)
        start_week = today - timedelta(days=today.weekday())
        end_week = start_week + timedelta(days=6)

        weekly_data = WeeklyResults.objects.filter(
            start=start_week.date(),
            end=end_week.date()
        ).last()
        weekly_result = eval(weekly_data.result) if weekly_data else []
        weekly_users = User.objects.filter(id__in=[p['player'] for p in weekly_result])
        context['weeks_winner'] = [
            {"user": weekly_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in weekly_result]
        winners = LeaderResults.objects.all().last()
        winners_data = eval(winners.result) if winners else []
        winners_users = User.objects.filter(id__in=[p['player'] for p in winners_data])
        context['winners'] = [
            {"user": winners_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in winners_data
        ]
        context['abc'] = ['A', 'B', 'C', 'D', 'E']
        context['start_time'] = ['1', '2', '3', '4', '5', '6', '7', '8', '9']
        try:
            context['count'] = GameTime.objects.filter(Q(player=self.request.user),
                                                       Q(player__unlimited=False),
                                                       Q(game_time__gte=day_range_utc()[0]),
                                                       Q(game_time__lte=day_range_utc()[1])).count()
            context['statistic'] = GameTime.objects.filter(player=self.request.user).order_by('-game_time').first()
        except:
            pass
        return context

    def post(self, request):
        if request.is_ajax():
            question_id = request.POST.get('question_id')
            answer_id = request.POST.get('answer_id')
            count = request.POST.get('count')
            game = GameTime.objects.filter(player=request.user,
                                           game_session=request.user.game_session
                                           ).order_by('-game_time').first()
            if game and not game.over:
                ans = Answer.objects.get(id=answer_id)
                try:
                    q = UserQuestionHistory.objects.get(
                        player_game=game,
                        question_id=question_id,
                        answer=ans.status
                    )
                    game.over = True
                    game.save()
                    return HttpResponse("False")
                except:
                    q = UserQuestionHistory(
                        player_game=game,
                        question_id=question_id,
                        answer=ans.status
                    )
                    q.save()
                try:
                    exclude_game = eval(game.player.exclude_question)
                    exclude_game.append(question_id)
                    game.player.exclude_question = str(exclude_game)
                    game.player.save()
                except:
                    pass
                if ans.status:
                    if int(count) == 10:
                        return JsonResponse({"last_question": "true"})
                    else:
                        return HttpResponse("True")
                else:
                    if game.two_question_joker and not game.two_question_joker_expired:
                        game.two_question_joker_expired = True
                        game.save()
                        return HttpResponse("False")
                    else:
                        game.over = True
                        game.save()
                        return HttpResponse("False")
            else:
                return JsonResponse({'token': 'error'})

    def check_user_social_share_permission(self):
        if self.request.user.social_share_count >= 1 and self.request.user.share_time != loc(timezone.now()).strftime(
                "%Y-%m-%d"):
            self.request.user.social_share_count = 0
            self.request.user.share_time = loc(timezone.now()).strftime("%Y-%m-%d")
            self.request.user.save()
            return True
        else:
            return False

    @method_decorator(login_required)
    def dispatch(self, *args, **kwargs):
        return super(OnlineTestView, self).dispatch(*args, **kwargs)


class ChangeQuestionsView(TemplateView):
    template_name = 'production/inline/question.html'

    def post(self, request, **kwargs):
        base_id = request.POST.get('q_id')
        exclude_list = eval(request.POST.get('q_list'))
        game = GameTime.objects.filter(player=request.user).order_by('-game_time').first()
        if not game.change_question_joker:
            game.change_question_joker = True
            game.save()
            if eval(request.user.exclude_question):
                old_exclude = list(set(eval(request.user.exclude_question)))
                res_exclude = exclude_list + old_exclude
                return render(request, self.template_name, self.get_context_data(res_exclude, **kwargs))
            else:
                return render(request, self.template_name, self.get_context_data(exclude_list, **kwargs))
        else:
            return JsonResponse({'error': 'Artıq bir dəfə istifadə olunub 2 ci dəfə istifadə qadağandır'})

    def get_context_data(self, base_id, **kwargs):
        context = super(ChangeQuestionsView, self).get_context_data(**kwargs)
        context['abc'] = ['A', 'B', 'C', 'D', 'E']
        context['count_base'] = self.request.POST.get('count')
        context['questions'] = Question.objects.all().order_by('?').exclude(pk__in=base_id)[:1]
        return context


class RightAnswerHalfView(View):
    def post(self, request):
        if request.is_ajax():
            question_id = request.POST.get('q_id')
            game = GameTime.objects.filter(player=request.user).order_by('-game_time').first()
            if not game.half_question_joker:
                game.half_question_joker = True
                game.save()
                question = Question.objects.get(id=question_id)
                result = question.get_true_answer()
                data_id = random.sample(result, len(result))
                return JsonResponse({'object': data_id})
            else:
                return JsonResponse({'error': 'Artıq bir dəfə istifadə olunub 2 ci dəfə istifadə qadağandır'})


class TwoAnswerJokerView(View):
    def post(self, request):
        if request.is_ajax():
            question_id = request.POST.get('q_id')
            game = GameTime.objects.filter(player=request.user).order_by('-game_time').first()
            if not game.two_question_joker:
                game.two_question_joker = True
                game.save()
                question = Question.objects.get(id=question_id)
                result = question.get_true_answer()
                data_id = random.sample(result, len(result))
                return JsonResponse({'object': data_id})
            else:
                return JsonResponse({'error': 'Artıq bir dəfə istifadə olunub 2 ci dəfə istifadə qadağandır'})


class AddTest(View):
    def get(self, request, *args, **kwargs):
        data = open(settings.BASE_DIR + "/test.txt", 'r')
        start = 0
        q_id_times = 0
        ques_id_list = []
        ra = 0
        ra_id = []
        for item in data.readlines():
            if start == 0 or start % 6 == 0:
                ques = item.split('\t')[1]
                q = Question(
                    question=ques
                )
                q.save()
                ques_id_list.append(q.id)
                if start == 0:
                    q_id_times += 0
                else:
                    q_id_times += 1
                    ra -= 5
            else:
                answ = item.split('\t')[1]
                a = Answer(
                    answer=answ,
                    question_id=ques_id_list[q_id_times]
                )
                a.save()
                ra_id.append(a.id)
                ra += 1
                # if ra == 5:
                #     d = random.randint(0,4)
                #     right = RightAnswer(
                #         answer_id=ra_id[d],
                #         question_id=ques_id_list[q_id_times]
                #     )
                #     right.save()
                #     ra_id[:] = []
            start += 1
        return HttpResponse("Cool")


def index(request):
    return HttpResponse("cool")


class UserResultsView(TemplateView):
    template_name = 'production/result/index.html'

    def get(self, request, *args, **kwargs):
        if request.GET.get('game_session'):
            game = GameTime.objects.filter(game_session=request.GET.get('game_session')).order_by('-game_time').first()
            if not game.done:
                c = loc(timezone.now()) - game.game_time
                duration = divmod(c.days * 86400 + c.seconds, 60)
                seconds = duration[1] - 3 if duration[1] > 3 else duration[1]
                game.duration = timezone.timedelta(minutes=duration[0], seconds=seconds)
                game.done = True
                game.save()
                # result = game.player.count_point()
                # game.player.point = result[0]
                # game.player.point_month = result[1]
                # game.player.point_week = result[2]
                # game.player.result_num = random.randint(1,9)
                # game.player.save()
                # leaderbord_count_week_and_month.delay(game.player.id)
                # try:
                #     CheckLeaderBorad(game)
                #     increase_point(game.player)
                # except:
                #     pass
                # count user points
        else:
            if request.user.is_authenticated:
                game = GameTime.objects.filter(player=request.user).order_by('-game_time').first()
                return redirect(reverse("base:result") + str('?game_session=') + str(game.game_session))
            else:
                return redirect('/')
        return render(request, self.template_name, self.get_context_data())

    def day_range(self):
        current_time = timezone.now()
        if current_time.hour >= 20:
            day_start = timezone.datetime(current_time.year,
                                          current_time.month,
                                          current_time.day,
                                          20,
                                          0).replace(tzinfo=current_time.tzinfo)
            day_end = timezone.datetime(current_time.year,
                                          current_time.month,
                                          current_time.day+1,
                                          20,
                                          0).replace(tzinfo=current_time.tzinfo)
        else:
            day_start = timezone.datetime(current_time.year,
                                          current_time.month,
                                          current_time.day-1,
                                          20,
                                          0).replace(tzinfo=current_time.tzinfo)
            day_end = timezone.datetime(current_time.year,
                                        current_time.month,
                                        current_time.day,
                                        20,
                                        0).replace(tzinfo=current_time.tzinfo)

        return [day_start, day_end]

    def get_context_data(self, **kwargs):
        context = {}
        context['abc'] = ['A', 'B', 'C', 'D', 'E']
        today = timezone.localtime(timezone.now())
        start_month = 1
        end_month = calendar.monthrange(today.year, today.month)[1]

        start_month = timezone.datetime.strptime("{0}-{1}-{2}".format(today.year, today.month, start_month), '%Y-%m-%d')
        start_month = start_month.replace(tzinfo=today.tzinfo)

        end_month = timezone.datetime.strptime("{0}-{1}-{2}".format(today.year, today.month, end_month), '%Y-%m-%d')
        end_month = end_month.replace(tzinfo=today.tzinfo)

        monthly_data = MonthlyResults.objects.filter(
            start=start_month.date(),
            end=end_month.date()
        ).last()
        monthly_result = eval(monthly_data.result) if monthly_data else []
        monthly_users = User.objects.filter(id__in=[p['player'] for p in monthly_result])
        context['months_winner'] = [
            {"user": monthly_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in monthly_result]

        # start_week = timezone.datetime(today.year, today.month - 1, 25)  # default
        # end_week = timezone.datetime(today.year, today.month, 8)
        start_week = today - timedelta(days=today.weekday())
        end_week = start_week + timedelta(days=6)

        weekly_data = WeeklyResults.objects.filter(
            start=start_week.date(),
            end=end_week.date()
        ).last()
        weekly_result = eval(weekly_data.result) if weekly_data else []
        weekly_users = User.objects.filter(id__in=[p['player'] for p in weekly_result])
        context['weeks_winner'] = [
            {"user": weekly_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in weekly_result]
        winners = LeaderResults.objects.all().last()
        winners_data = eval(winners.result) if winners else []
        winners_users = User.objects.filter(id__in=[p['player'] for p in winners_data])
        context['winners'] = [
            {"user": winners_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in winners_data
        ]
        try:
            context['count'] = GameTime.objects.filter(Q(player=self.request.user),
                                                       Q(player__unlimited=False),
                                                       Q(game_time__gte=day_range_utc()[0]),
                                                       Q(game_time__lte=day_range_utc()[1])).count()
        except:
            pass
        try:
            context['statistic'] = GameTime.objects.filter(game_session=self.request.GET.get('game_session')).order_by('-game_time').first()
        except:
            pass
        context['redirect'] = self.request.GET.get('index',False)
        return context


class UserGameHistory(TemplateView):
    template_name = 'production/history/index.html'

    def get_context_data(self, **kwargs):
        context = {}
        today = timezone.localtime(timezone.now())
        start_month = 1
        end_month = calendar.monthrange(today.year, today.month)[1]

        start_month = timezone.datetime.strptime("{0}-{1}-{2}".format(today.year, today.month, start_month), '%Y-%m-%d')
        start_month = start_month.replace(tzinfo=today.tzinfo)

        end_month = timezone.datetime.strptime("{0}-{1}-{2}".format(today.year, today.month, end_month), '%Y-%m-%d')
        end_month = end_month.replace(tzinfo=today.tzinfo)

        monthly_data = MonthlyResults.objects.filter(
            start=start_month.date(),
            end=end_month.date()
        ).last()
        monthly_result = eval(monthly_data.result) if monthly_data else []
        monthly_users = User.objects.filter(id__in=[p['player'] for p in monthly_result])
        context['months_winner'] = [
            {"user": monthly_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in monthly_result]

        # start_week = timezone.datetime(today.year, today.month - 1, 25)  # default
        # end_week = timezone.datetime(today.year, today.month, 8)
        start_week = today - timedelta(days=today.weekday())
        end_week = start_week + timedelta(days=6)

        weekly_data = WeeklyResults.objects.filter(
            start=start_week.date(),
            end=end_week.date()
        ).last()
        weekly_result = eval(weekly_data.result) if weekly_data else []
        weekly_users = User.objects.filter(id__in=[p['player'] for p in weekly_result])
        context['weeks_winner'] = [
            {"user": weekly_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in weekly_result]
        winners = LeaderResults.objects.all().last()
        winners_data = eval(winners.result) if winners else []
        winners_users = User.objects.filter(id__in=[p['player'] for p in winners_data])
        context['winners'] = [
            {"user": winners_users.get(id=x['player']), "duration": x['duration'], "answer_count": x['answer_count']}
            for x in winners_data
        ]
        context['calendars'] = get_four_month_calendar()
        context['games'] = GameTime.objects.filter(player=self.request.user, done=True)
        return context


class AboutUsView(TemplateView):
    template_name = 'development/about/index.html'


class ContactView(TemplateView):
    template_name = 'development/contact/index.html'


class GameExitView(View):
    def get(self, request, *args, **kwargs):
        game = GameTime.objects.filter(player=request.user).order_by('-game_time').first()
        if not game.done:
            c = loc(timezone.now()) - game.game_time
            duration = divmod(c.days * 86400 + c.seconds, 60)
            game.duration = timezone.timedelta(minutes=duration[0], seconds=duration[1])
            game.done = True
            game.save()
        return redirect("/")


#
class ShareFacebookWithFriends(View):
    share_url = 'http://www.facebook.com/dialog/send'
    mobile_url = 'https://www.facebook.com/dialog/share'
    redirect_url = 'http://mok25.az/social/results/'
    share_link = 'http://mok25.az/'

    def get(self, request):
        if request.GET.get('mobile', False):
            params = dict(app_id=settings.SOCIAL_AUTH_FACEBOOK_KEY,
                          display="popup",
                          href=self.share_link,
                          redirect_uri=self.redirect_url)
            encode_params = self.params_url_encode(params)
            return redirect(self.mobile_url + encode_params)
        else:
            params = dict(app_id=settings.SOCIAL_AUTH_FACEBOOK_KEY,
                          redirect_uri=self.redirect_url,
                          link=self.share_link)
            encode_params = self.params_url_encode(params)
            return redirect(self.share_url + encode_params)

    def params_url_encode(self, item):
        encode_url = ""
        for key, value in item.items():
            encode_url += str(key) + "=" + str(value) + "&"
        encode_url = "?%s" % encode_url[:-1]
        return encode_url

    @method_decorator(login_required)
    def dispatch(self, *args, **kwargs):
        return super(ShareFacebookWithFriends, self).dispatch(*args, **kwargs)


class SocialResultsView(View):
    def get(self, request, *args, **kwargs):
        if request.GET.get('success', False):
            share_count = int(request.GET.get('success'))
            user = User.objects.get(id=request.user.id)
            user.social_share_count += share_count
            user.save()
            if user.social_share_count >= 1:
                return redirect(reverse("base:player-dashboard"))
            else:
                return redirect(reverse("base:player-dashboard"))
        else:
            share_count = 0
            user = User.objects.get(id=request.user.id)
            user.social_share_count = 0
            user.save()
            return redirect(reverse("base:player-dashboard"))

    @method_decorator(login_required)
    def dispatch(self, *args, **kwargs):
        return super(SocialResultsView, self).dispatch(*args, **kwargs)


# Error handling
def get_error_page(request):
    return redirect("/")

# share result page view
class SharePlayerResultView(TemplateView):
    template_name = 'production/share/index.html'

    def get(self, request, *args, **kwargs):
        if request.GET.get('g_s',False):
            game_session = request.GET.get('g_s')
            return render(request, self.template_name, self.get_context_data(game_session))
        else:
            return redirect("/")

    def get_context_data(self, game_session):
        context = super(SharePlayerResultView, self).get_context_data()
        if game_session:
            context['statistic'] = GameTime.objects.filter(game_session=game_session).order_by('-game_time').first()
            context['header'] = WebsiteHeader.objects.all().last()
        return context


class CeleryWorkerView(View):
    def get(self, request, *args, **kwargs):
        if request.user.is_superuser:
            calc_weekly()
            calc_monthly()
            calc_leaderbord()
            return JsonResponse({"week_leaderbord":"updated", "month_leaderbord":"updated", "all_leaderbord":"updated"})
        else:
            return JsonResponse({"error":"invalid_client","error_description":"Invalid client"})


class PlayerDashboard(OnlineTestView):
    template_name = 'production/dashboard/index.html'
    inline_template = 'production/inline/friends.html'

    def get(self, request, *args, **kwargs):
        fb_data = FacebookFriendsList(self.request.user)
        data = fb_data.get_user_friends()
        if data:
            users = UserSocialAuth.objects.filter(uid__in=[u['id'] for u in data]).order_by("-user__point")
        else:
            users = None
        return render(request, self.template_name, self.get_context_data(users))

    def post(self, request, *args, **kwargs):
        count = request.POST.get('count', False)
        if count:
            fb_data = FacebookFriendsList(self.request.user)
            data = fb_data.get_user_friends()
            if data:
                users = UserSocialAuth.objects.filter(uid__in=[u['id'] for u in data]).order_by("-user__point")
            else:
                users = None

            return render(request, self.inline_template, self.get_context_data(users))
        else:
            return JsonResponse({'error':'Count not found'})

    def get_context_data(self, data=None, **kwargs):
        context = super(PlayerDashboard, self).get_context_data(**kwargs)
        context['count'] = GameTime.objects.filter(Q(player=self.request.user),
                                                   Q(player__unlimited=False),
                                                   Q(game_time__gte=day_range_utc()[0]),
                                                   Q(game_time__lte=day_range_utc()[1])).count()
        if data:
            paginator = Paginator(data, 10)
            if self.request.POST.get('count', False):
                page = self.request.POST.get('count')
            else:
                page = 1
            try:
                contents = paginator.page(page)
            except PageNotAnInteger:
                contents = paginator.page(1)
            except EmptyPage:
                contents = paginator.page(paginator.num_pages)
            context['data'] = contents
            context['pages'] = paginator
        else:
            context['data'] = None
        return context


class ExportAllQuestions(View):

    def get(self, request, *args, **kwargs):
        if request.user.is_superuser:
            question_list = Question.objects.all()
            paginator = Paginator(question_list, 500)
            try:
                if request.GET.get('page',False):
                    content = paginator.page(int(request.GET.get('page')))
                else:
                    content = paginator.page(1)
            except:
                content = paginator.page(1)
            document = Document()
            document.add_heading('Mok25 Suallar', 0)
            document.add_heading('Sualların sayı: {}'.format(question_list.count()), level=6)
            for question in content:
                document.add_paragraph(
                    question.question,
                    style='List Number'
                )
                if question.image != '':
                    image_path = os.path.join(settings.MEDIA_ROOT, question.image.name)
                    document.add_paragraph('http://mok25.az/media/'+question.image.name)
                table = document.add_table(rows=2, cols=2)
                hdr_cells = table.rows[0].cells
                second = table.rows[1].cells
                answer_list = question.answer_set.all()
                hdr_cells[0].text = 'A) {}{}'.format(answer_list[0].answer,"(*)" if answer_list[0].status else "")
                hdr_cells[1].text = 'B) {}{}'.format(answer_list[1].answer,"(*)" if answer_list[1].status else "")
                second[0].text = 'C) {}{}'.format(answer_list[2].answer,"(*)" if answer_list[2].status else "")
                second[1].text = 'D) {}{}'.format(answer_list[3].answer,"(*)" if answer_list[3].status else "")
                document.add_paragraph('')
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=suallar_page{}.docx'.format(self.request.GET.get('page','1'))
            document.save(response)
            return response
        else:
            return JsonResponse({"error": "invalid_client", "error_description": "Invalid client"})


class ListViewAllDocuments(View):

    def get(self, request, *args, **kwargs):
        if request.user.is_superuser:
            documents = Question.objects.all().count()
            result = []
            ul = "<ul>"
            for doc in range(math.ceil(documents/500)):
                result.append("<li><a href='/export/document?page={0}'>suallar_page{0}.docx</a></li>".format(doc+1))
            output = ul + "".join(result) + "</ul>"
            return HttpResponse(output)
        else:
            return JsonResponse({"error": "invalid_client", "error_description": "Invalid client"})